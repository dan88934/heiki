from flask import render_template, request, escape, Flask, flash, redirect, url_for, Blueprint, send_from_directory
from werkzeug.utils import secure_filename
import os
from uuid import uuid4
import pytesseract
# import pykakasi
from sudachipy import tokenizer, dictionary
from JPVL_jamdict import Jamdict # removed .
from docx import Document
from docx.shared import Pt
import docx2txt
import re
import sys
try:
    from PIL import Image
except ImportError:
    import Image
from zipfile import BadZipFile
    

UPLOAD_FOLDER = '/uploaded_files'
ALLOWED_EXTENSIONS = {'docx', 'jpg', 'jpeg','png'}

original_words = [] #Words for column 1 will be put in here
kana_and_eng_def = [] #Words for column 2 will be put in here
combined = [] #The loop will place all the words from each of the images passed in (in the case of multiple images) in here when it concatinates the list of words for each image together.

#1. ===Define functions===

def allowed_file(filename): #Checks that file matches on in the set of allowed_extensions above
    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def make_unique(string): #This adds 4 unique characters on the end of the string (to avoid duplicate files being created)
    ident = uuid4().__str__()[:4]
    return f"{string}-{ident}"

def get_japanese_only(context): #Here we remove all non-Japanese (excluding Chinese) characters from the string
    filtrate = re.compile(u'[^\u4E00-\u9FA5\u3040-\u309f\u30A0-\u30FF]')  #These are the unicode ranges for CJK, Hiragana and Katanaka (all characters that do not fall into these groups are excluded)
    context = filtrate.sub(r'', context)
    return context

def tokenize_into_words(string): #Takes a string of Japanese-only text and turns it into a list of words (and some irrelevant hiragana and karakana we will remove later)
    mode = tokenizer.Tokenizer.SplitMode.A
    tokenizer_obj = dictionary.Dictionary().create()
    tokenized_list = [m.dictionary_form() for m in tokenizer_obj.tokenize(string, mode)]
    return tokenized_list

def list_concat(master_list): # This turns our multiple lists (one for each image or document) into a single list which can then be iterated over
    for item in master_list:
        combined.append(item)
    return combined

def remove_duplicate_items(input_list): #This removes duplicate items from our list by turning it into a dictionary and then converting it back again
    duplicates_removed_list = list(
            dict.fromkeys(input_list))
    return duplicates_removed_list

def remove_irrelevant_characters(input_list): #This removes some verb components and any random kanji which tend to occur and have no meaning on their own (e.g. ます)
    unwanted_particles_etc = {'ます','引', 'って'}  #引 appears in the list sometimes when it is not actually in the image
    unwanted_particles_removed_list = [item for item in input_list if item not in unwanted_particles_etc] #We use a list comprehension to remove them
    return unwanted_particles_removed_list

def remove_single_katakana(input_list): #Single katanaka characters are removed from the list (they do not have a meaning on their own)
    single_katakana = {'ア','イ','ウ','エ','オ','カ','キ','く','ケ','コ','サ','シ','ス','セ','ソ','タ','チ','ツ','テ','ト','ナ','ニ','ヌ','ネ',
    'ノ','ハ','ヒ','フ','ヘ','ホ','マ','ミ','ム','メ','モ','ヤ','ユ','ヨ','ラ','リ','ル','レ','ロ','ワ','ヲ'} #All commonly used katakana
    single_katanaka_removed_list = [item for item in input_list if item not in single_katakana]
    return single_katanaka_removed_list

def remove_single_hiragana(input_list): #Single hiragana characters are removed from the list (they do not have a meaning on their own)
    single_hiragana = {'あ','い','う','え','お','か','き','く','け','こ','が','ぎ','ぐ','げ','ご','さ','し','す','せ','そ','ざ','じ',
    'ず','ぜ','ぞ','た','ち','つ','て','と','だ','ぢ','づ','で','ど','な','に','ぬ','ね','の','は','ひ','ふ','へ','ほ','ば','び','ぶ',
    'べ','ぼ','ぱ','ぴ','ぷ','ぺ','ぽ','ま','み','む','め','も','や','ゆ','よ','ら','り','る','れ','ろ','わ','を','ん'}# All commonly used Hiragana
    single_hiragana_removed_list = [item for item in input_list if item not in single_hiragana]
    return single_hiragana_removed_list

def add_items_to_original_word_list(input_list): #Appends words to the original words list (words which will go in column 1) 
    if not input_list: #Error handling for the event that input_list is empty (this would occur in the event that there were not Japanese characters in the image or document)
        print('Error - No Japanese words in images')
        sys.exit(1)
    for item in input_list:
        original_words.append(item)
    return original_words

# @measure #Decorator to measure how long a function takes to run
def get_reading_and_eng(input_list): #This uses Jamdict to return a list of strings, each string contains the hiragana and english definitions for each of the words in our original words list (above function)
    jmd = Jamdict()
    for item in input_list:   
        result = jmd.lookup(item)
        kana_and_eng_def.append(str(result))
    return kana_and_eng_def


process_file = Blueprint('process_file', __name__)
@process_file.route('/', methods=['GET', 'POST']) #This is the homepage which will receive the input
def index():
    ###########
    #Input    #
    ###########
    if request.method == 'POST': #Arguably it would be nice to have some input error handling/validation here
        global title_w_unique_id
        title_w_unique_id = []
        global input_files
        input_files = [] #This contains the names of each file to be processed
        global input_title
        input_title = request.form['title']
        global input_type
        input_type = request.form['file-type']
        files = request.files.getlist("file")
        for file in files: #Loops through each file uploaded, saves them in uploaded_files folder and appends the filename to input_files
            if file and allowed_file(file.filename):
                original_filename = secure_filename(input_title)
                unique_filename = make_unique(original_filename) 
                title_w_unique_id.append(unique_filename) #Append title + unique number
                file.save(os.path.join('./uploaded_files/', unique_filename))
                input_files.append(url_for('process_file.uploaded_file', filename=unique_filename))
        #################
        #File Processing#
        #################
        # pytesseract.pytesseract.tesseract_cmd = r'/home/heiki/.linuxbrew/bin/tesseract' # PythonAnywhere version
        pytesseract.pytesseract.tesseract_cmd = r'/usr/local/Cellar/tesseract/5.0.0/bin/tesseract' #15 MBP development version
        tokenized_list = []
        combined_japanese_chars_list = []
        file_type = input_type #This checks the file type entered on the web app dropdown box (doc or img) and sends the files to the correct processing condition
        if file_type == '': #This is actually not possible if user is entering data via the gui on the webpage (There is no blank choice)
            print('Error - Please select a file type') #Error handling for the event that no file-type is selected
            sys.exit(1)
        for item in input_files:
            file_location = './uploaded_files' + item #This gives it the correct path
            if file_type == 'document':
                try:
                    raw_text = docx2txt.process(file_location) #Get a string of all the text in the doc
                    print(raw_text)
                except FileNotFoundError:
                    print('Error - Document does not exist on server (It may exist on client side)')
                    sys.exit(1)
                except BadZipFile: # Occurs when image is sent and the document option is selected
                    print('Error - Posted file is not a document')
                    sys.exit(1)
                japanese_only = get_japanese_only(raw_text) #Remove non-Japanese text from the string
                print(japanese_only)
                tokenized_list = tokenize_into_words(japanese_only) #Tokenize string into a list of words
                print(tokenized_list)
                combined_japanese_chars_list = list_concat(tokenized_list) #join the lists created for each file passed in together
                print(combined_japanese_chars_list)
            elif file_type == 'image':
                try:
                    raw_text = pytesseract.image_to_string( #Get a string of all the text in the image via OCR
                        Image.open(file_location), lang='jpn')
                    print(raw_text)
                except FileNotFoundError:
                    print('Error - Image does not exist')
                    sys.exit(1)
                japanese_only = get_japanese_only(raw_text) #Remove non-Japanese text from the string
                tokenized_list = tokenize_into_words(japanese_only) #Tokenize string into a list of words
                combined_japanese_chars_list = list_concat(tokenized_list) #join the lists created for each file passed in together
        #Exit loop and back to process_file() scope
        input_files.clear()
        irrelevant_items_removed_list = remove_irrelevant_characters(combined_japanese_chars_list) #Remove items that have no meaning (see function for more details)
        single_katakana_removed_list = remove_single_katakana(irrelevant_items_removed_list) #Remove single katanaka characters (which are meaningless on their own)
        single_hiragana_removed_list = remove_single_hiragana(single_katakana_removed_list) #Removes single hiragana characters (which are meaningless on their own)
        duplicate_items_removed_list = remove_duplicate_items(single_hiragana_removed_list) #Removes any duplicate words and characters in our word list
        add_items_to_original_word_list(duplicate_items_removed_list) #Add original words to list 
        get_reading_and_eng(duplicate_items_removed_list) #Add reading and eng definition to list 
        print(kana_and_eng_def)
        #===clearing data for next user===
        tokenized_list.clear()
        combined_japanese_chars_list.clear()
        irrelevant_items_removed_list.clear()
        single_katakana_removed_list.clear()
        single_hiragana_removed_list.clear()
        duplicate_items_removed_list.clear()

        #################
        #Create Document#
        #################

        print('Entering create vocab list')
        name = input_title #This is the title and file name which was set in f_input. It was nessesary to append the string to a list
        document = Document()
        document.add_heading(name, 0) #Sets heading of document to value set in title


        #Make a tuple of tuples
        word_list = tuple(zip(original_words, kana_and_eng_def))
        original_words.clear()
        kana_and_eng_def.clear()
        #Create table
        table = document.add_table(rows=1, cols=2) #Row number will increase automatically as content is looped in
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Word' #This and below line set the title for each of the two columns
        hdr_cells[1].text = 'Reading + Meanings'
        hdr_cells[0].width = 1097280 #This and below line define how wide the columns are
        hdr_cells[1].width = 48463200

        for word, meaning in word_list:
            #Here we loop through the tuple of tuples of word and meanings and add them to the table
            row_cells = table.add_row().cells
            row_cells[0].text = word
            row_cells[1].text = meaning


        for row in table.rows:
            #This changes the font size of the table to 9 (reduces the number of pages)
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(9)

        del word_list #Removing data from the tuple for the next user


        #Below we set the file name of the vocabulary list (if title box empty, default file name is 'vocabularly list - note that the frontend Bootstrap JS does not actually allow there to be no title)
        output_file_extension = '.docx'
        file_name = title_w_unique_id[0] + output_file_extension
        upload_directory = './output_files/'
        if not name:
            document.save('./output_files/' + 'Vocabulary list.docx')
            file_name = 'Vocabulary list.docx'
        else:
            document.save('./output_files/' + file_name)
        #Below we clean up by removing the original file (img or doc) from the uploaded_files directory
        for item in title_w_unique_id:
            uploaded_file_path = './uploaded_files/' + item #Path of the initially uploaded file (doc or image)
            os.remove(uploaded_file_path) #Removes the initially uploaded file (to save resources)
        title_w_unique_id.clear()
        render_template("oldindex.html")
        return send_from_directory(directory=upload_directory, path=file_name, as_attachment=True) #as_attachment keeps the original file name (the title set by user), rather than calling the file the name of the page
    return render_template("index.html")


@process_file.route('/<filename>')
def uploaded_file(filename):
    return send_from_directory('./uploaded_files/',
                               filename)